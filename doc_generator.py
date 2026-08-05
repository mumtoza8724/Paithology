from docx import Document
from docx.shared import Inches


def create_doc(
        medical_center,
        patient_name,
        research_number,
        admission_date,
        gender,
        birth_date,
        operation_type,
        organ,
        clinical,
        doctor,
        slides,
        cassettes,
        stain_method,
        lab_worker,
        result_date,
        material_color,
        material_size,
        material_consistency,
        image_path,
        description,
        conclusion
):

    doc=Document()

    doc.add_heading(
        "DDS PATHOLOGY",
        level=1
    )

    doc.add_heading(
        "Патогистологическое исследование",
        level=2
    )


    fields=[

        ("№ исследования",research_number),
        ("Дата поступления",admission_date),
        ("Направившее медицинское учреждение", medical_center),
        ("ФИО пациента",patient_name),
        ("Пол",gender),
        ("Дата рождения",birth_date),
        ("Вид операции",operation_type),
        ("Орган",organ),
        ("Клинический диагноз",clinical),
        ("Лечащий врач",doctor),
        ("Количество стекол",slides),
        ("Количество кассет",cassettes),
        ("Методы окраски",stain_method),
        ("Лаборант",lab_worker),
        ("Дата выдачи результата",result_date),
        ("Цвет материала",material_color),
        ("Размер материала",material_size),
        ("Консистенция",material_consistency)

    ]

    table = doc.add_table(
        rows=0,
        cols=2
    )

    table.style = "Table Grid"

    for title, value in fields:
        row_cells = (
            table.add_row().cells
        )

        row_cells[0].text = str(
            title
        )

        row_cells[1].text = str(
            value
        )

    doc.add_heading(
        "QR-код исследования",
        level=2
    )

    try:

        doc.add_picture(
            "static/qr.png",
            width=Inches(1.5)
        )

    except:
        pass

    doc.add_heading(
        "Фото микропрепарата",
        level=2
    )


    try:

        doc.add_picture(
            image_path,
            width=Inches(4)
        )

    except:
        pass


    doc.add_heading(
        "Микроскопическое описание",
        level=2
    )

    doc.add_paragraph(
        description
    )


    doc.add_heading(
        "Заключение",
        level=2
    )

    doc.add_paragraph(
        conclusion
    )


    footer=(
        doc
        .sections[0]
        .footer
    )

    footer.paragraphs[0].text=(

    "Примечание: окончательная постановка "
    "диагноза осуществляется лечащим врачом "
    "на основании клинических, лабораторных "
    "и инструментальных исследований, "
    "в связи с чем гистологические и "
    "цитологические заключения не являются "
    "окончательным диагнозом."

    )


    doc.save(
        "static/report.docx"
    )